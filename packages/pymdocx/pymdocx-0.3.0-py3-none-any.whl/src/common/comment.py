import docx


def create_comment_doc(doc_file_path, doc_file_path_new):
    # 创建一个有评论的doc
    document = docx.Document(docx=doc_file_path)
    # for p in document.paragraphs:
    #     print(p)
    #     comment = p.add_comment('测试', author='Obay Daba', initials='od')
    document.save(doc_file_path_new)


def has_comment(p):
    return 1 if p.comments else 0


def get_comment(doc_file_path):
    document = docx.Document(docx=doc_file_path)
    for p in document.paragraphs:
        for c in p.comments:
            print(c.text)


def get_comment_xml(doc_file_path):
    document = docx.Document(docx=doc_file_path)
    comment_xml = document.comments_part.element.xml
    return comment_xml


def parse_p_comment(p_obj):
    # docx.text.comment.Comment
    comment_info_list = []
    for comment in p_obj.comments:
        comment_info_list.append({
            'comment_text': comment.text,
            'dtime': comment.element.date,
            'author': comment.element.author,
            'initials': comment.element.initials,
            'id': comment.element._id,
        })
    return comment_info_list